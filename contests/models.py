import importlib
import io
import json
import os
import random
import zipfile

import docx
from django.contrib.auth.models import User
from django.contrib.contenttypes.fields import GenericForeignKey, GenericRelation
from django.contrib.contenttypes.models import ContentType
from django.core.exceptions import ObjectDoesNotExist
from django.core.validators import MaxValueValidator, MinValueValidator, validate_comma_separated_integer_list
from django.db import models
from django.db.models import Q
from django.dispatch import receiver
from django.urls import reverse
from django.utils import timezone

from accounts.models import Account, Comment, Faculty, Notification
from contest.abstract import CDEntry, CRDEntry, CRUDEntry
from contest.soft_deletion import SoftDeletionManager, SoftDeletionModel, SoftDeletionQuerySet
from contest.utils import transliterate

try:
    from tools.sandbox import get_sandbox_class
    from tools.utility import Status, diff
except ImportError:
    from contest.utils import Status, diff, get_sandbox_class

"""=================================================== Attachment ==================================================="""


def attachment_path(instance, filename):
    return "attachments/{app_label}/{model}/{id}/{filename}".format(app_label=instance.object._meta.app_label.lower(),
                                                                    model=instance.object._meta.object_name.lower(),
                                                                    id=instance.object.id,
                                                                    filename=filename)


class Attachment(CDEntry):
    object_type = models.ForeignKey(ContentType, on_delete=models.CASCADE)
    object_id = models.PositiveIntegerField()
    object = GenericForeignKey(ct_field='object_type')

    file = models.FileField(upload_to=attachment_path, verbose_name="Файл")

    def extension(self):
        return os.path.splitext(self.file.name)[1]

    class Meta(CDEntry.Meta):
        verbose_name = "Файл"
        verbose_name_plural = "Файлы"

    @property
    def dirname(self):
        return os.path.split(self.file.path)[0]

    @property
    def filename(self):
        return os.path.split(self.file.name)[1]

    def __str__(self):
        return self.filename


@receiver(models.signals.post_delete, sender=Attachment)
def auto_delete_file_on_delete(sender, instance, **kwargs):
    """ deletes file from filesystem when corresponding `Attachment` object is deleted. """
    if instance.file:
        if os.path.isfile(instance.file.path):
            os.remove(instance.file.path)


"""===================================================== Course ====================================================="""


class CourseQuerySet(SoftDeletionQuerySet):
    def of_faculty(self, faculty):
        if faculty.is_interfaculty:
            return self.filter(faculty_id=faculty.id)
        else:
            return self.filter(Q(faculty_id=faculty.id) | Q(faculty__short_name="МФК")).distinct()


class CourseManager(SoftDeletionManager):
    def get_queryset_for_contest(self, course, user):
        queryset = self.get_queryset()
        return queryset.filter(Q(id=course.id) | Q(faculty=user.account.faculty, leaders__id=user.id)).distinct()


class Course(SoftDeletionModel, CRUDEntry):
    LEVEL_CHOICES = Account.LEVEL_CHOICES

    faculty = models.ForeignKey(Faculty, on_delete=models.DO_NOTHING, verbose_name="Факультет")
    leaders = models.ManyToManyField(User, through='CourseLeader', through_fields=('course', 'leader'),
                                     related_name='courses_leading', verbose_name="Ведущие преподаватели")

    title_official = models.CharField(max_length=100, verbose_name="Официальное название",
                                      help_text="Официальное название будет использовано при составлении отчетных "
                                                "документов")
    title_unofficial = models.CharField(max_length=100, null=True, blank=True, verbose_name="Неофициальное название",
                                        help_text="Неофициальное название будет отображено для пользователей сайта")
    description = models.TextField(verbose_name="Описание", blank=True)
    level = models.PositiveSmallIntegerField(choices=LEVEL_CHOICES, verbose_name="Уровень")

    comment_set = GenericRelation(Comment, content_type_field='object_type')

    objects = CourseManager.from_queryset(CourseQuerySet)()

    class Meta(CRUDEntry.Meta):
        ordering = ('level', 'id')
        verbose_name = "Курс"
        verbose_name_plural = "Курсы"

    @property
    def title(self):
        return self.title_official if self.title_unofficial is None else self.title_unofficial

    def get_latest_submissions(self):
        return Submission.objects.filter(assignment__isnull=False, problem__contest__course=self)

    def get_discussion_url(self):
        return reverse('contests:course-discussion', kwargs={'pk': self.pk})

    def save(self, *args, **kwargs):
        created = self._state.adding
        super().save(*args, **kwargs)
        if created:
            Filter.objects.get_or_create(user=self.owner, course=self)

    def __str__(self):
        return self.title


"""================================================== CourseLeader =================================================="""


class CourseLeader(models.Model):
    GROUP_CHOICES = ((0, "Все"),) + Account.GROUP_CHOICES
    GROUP_DEFAULT = 0

    course = models.ForeignKey(Course, on_delete=models.CASCADE, related_name='+', verbose_name="Преподаваемый курс")
    leader = models.ForeignKey(User, on_delete=models.CASCADE, related_name='+', verbose_name="Преподаватель курса")

    group = models.PositiveSmallIntegerField(choices=GROUP_CHOICES, default=GROUP_DEFAULT, verbose_name="Группа")
    subgroup = models.PositiveSmallIntegerField(choices=GROUP_CHOICES, default=GROUP_DEFAULT, verbose_name="Подгруппа")

    class Meta:
        constraints = [models.UniqueConstraint(fields=('course', 'leader'), name='unique_course_leader')]
        verbose_name = "Преподаватель курса"
        verbose_name_plural = "Преподаватели курсов"

    def __str__(self):
        return f"{self.leader.account.get_short_name()} преподает {self.course}"

    def save(self, *args, **kwargs):
        created = self._state.adding
        super().save(*args, **kwargs)
        if created:
            Filter.objects.get_or_create(user=self.leader, course=self.course)


"""===================================================== Credit ====================================================="""


def generate_credit_report(group_name, students, report_type, examiners, faculty, direction, discipline, semester,
                           date):
    months = ['января', 'февраля', 'марта', 'апреля', 'мая', 'июня', 'июля', 'августа', 'сентября', 'октября', 'ноября',
              'декабря']

    f = open('contest/documents/blank_report.docx', 'rb')

    document = docx.Document(f)
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = docx.shared.Pt(10)

    # -------------------------------------------------------------------------

    header_paragraph = document.paragraphs[0]
    header_paragraph.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    header_paragraph.add_run('''ФИЛИАЛ МОСКОВСКОГО ГОСУДАРСТВЕННОГО УНИВЕРСИТЕТА
имени М.В.ЛОМОНОСОВА в г. ТАШКЕНТЕ
''')

    run = header_paragraph.add_run('ЭКЗАМЕНАЦИОННАЯ')
    run.bold = True
    run.underline = True

    header_paragraph.add_run(' (ЗАЧЕТНАЯ) ВЕДОМОСТЬ №\n')

    # -------------------------------------------------------------------------

    group_paragraph = document.paragraphs[1]

    group_paragraph.add_run('ФАКУЛЬТЕТ  ')
    group_paragraph.add_run(faculty).bold = True
    group_paragraph.add_run('    НАПРАВЛЕНИЕ  ')
    group_paragraph.add_run(direction).bold = True
    group_paragraph.add_run('    ГРУППА  ')
    group_paragraph.add_run(group_name).bold = True
    group_paragraph.add_run('    СЕМЕСТР  ')
    group_paragraph.add_run(str(semester) + '\n').bold = True

    # -------------------------------------------------------------------------

    date_paragraph = document.paragraphs[2]

    date_paragraph.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    formatted_date = "\nДАТА  «{}» {} {} г.\n".format(date.day, months[date.month - 1], date.year)
    date_paragraph.add_run(formatted_date).bold = True

    # -------------------------------------------------------------------------

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

    paragraph.add_run('''ВСЕГО ОЦЕНОК_______
ОТЛИЧНО_______
ХОРОШО_______
УДОВЛЕТВОРИТЕЛЬНО_______
НЕУДОВЛЕТВОРИТЕЛЬНО_______
ЗАЧЕТ_______
НЕЗАЧЕТ_______
НЕ ЯВИЛСЯ_______''')

    # -------------------------------------------------------------------------

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

    paragraph.add_run("ДЕКАН ФАКУЛЬТЕТА\n\nЗАВЕДУЮЩИЙ КАФЕДРЫ\n\nПОДПИСИ ЭКЗАМЕНАТОРОВ")

    # -------------------------------------------------------------------------

    credit_info_table = document.tables[0]
    credit_info_table.rows[0].cells[1].text = discipline
    credit_info_table.rows[1].cells[1].text = report_type.upper()
    credit_info_table.rows[2].cells[1].text = ' / '.join(examiners)

    # -------------------------------------------------------------------------

    credit_info_table = document.tables[1]

    for number, student in enumerate(students, 1):
        row = credit_info_table.add_row()
        row.height_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = docx.shared.Cm(0.635)

        row.cells[0].text = str(number)

        row.cells[1].text = student["name"]
        row.cells[1].paragraphs[0].runs[0].font.size = docx.shared.Pt(12)

        row.cells[2].text = str(student["record_book_id"])
        row.cells[2].paragraphs[0].runs[0].font.size = docx.shared.Pt(12)

        row.cells[3].text = str(student["score"])
        row.cells[3].paragraphs[0].runs[0].font.size = docx.shared.Pt(12)

    # -------------------------------------------------------------------------

    target_stream = io.BytesIO()
    document.save(target_stream)
    f.close()

    return target_stream.getvalue()


class CreditManager(models.Manager):
    def create_set(self, owner, course, accounts):
        user_ids = accounts.values_list('user_id', flat=True)
        new_credits = []
        for user_id in user_ids:
            new_credit = Credit(owner_id=owner.id, user_id=user_id, course_id=course.id)
            new_credits.append(new_credit)
        return self.bulk_create(new_credits)

    def create_report(self, course, group_name, students, type, examiners, faculty, discipline, semester, date):
        examiners = examiners.annotate(
            lead=models.Exists(course.leaders.filter(id=models.OuterRef('user_id')))
        ).order_by('-lead')
        students_with_scores = students.with_credits(course)

        students_prepared = []
        for student in students_with_scores:
            score_choices = {
                0: "",
                2: "2 (неуд.)",
                3: "3 (удов.)",
                4: "4 (хор.)",
                5: "5 (отл.)",
            }
            students_prepared.append({
                "name": str(student),
                "score": score_choices[student.credit_score],
                "record_book_id": student.record_book_id or "",
            })

        examiners = [(examiner.position + " " if examiner.position else "") + str(examiner) for examiner in examiners]
        report_file = generate_credit_report(group_name=group_name, students=students_prepared, report_type=type,
                                             examiners=examiners, faculty=faculty, direction=faculty,
                                             discipline=discipline, semester=semester, date=date)

        filename = "vedomost_{}_{}_{}_{}".format(type.lower(), course.title.lower(), group_name.lower(), date)
        filename = transliterate(filename).replace(" ", "_")

        return report_file, filename


class Credit(CRUDEntry):
    SCORE_CHOICES = (
        (5, "отлично"),
        (4, "хорошо"),
        (3, "удовлетворительно"),
        (2, "неудовлетворительно"),
        (0, "нет оценки"),
    )
    DEFAULT_SCORE = 0

    owner = models.ForeignKey(User, on_delete=models.CASCADE, related_name='+', verbose_name="Владелец")
    user = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name="Студент")
    course = models.ForeignKey(Course, on_delete=models.CASCADE, verbose_name="Курс")

    score = models.PositiveSmallIntegerField(choices=SCORE_CHOICES, default=DEFAULT_SCORE, verbose_name="Оценка")

    objects = CreditManager()

    class Meta(CRUDEntry.Meta):
        unique_together = ('user', 'course')
        permissions = [
            ("report_credit", "Генерировать ведомость"),
        ]
        ordering = ('-course',)
        verbose_name = "Зачет"
        verbose_name_plural = "Зачеты"

    def __str__(self):
        return f"Зачет {self.user.account.get_short_name()} по курсу: {self.course}"


"""=================================================== Attendance ==================================================="""


class Attendance(CDEntry):
    owner = models.ForeignKey(User, on_delete=models.CASCADE, related_name='+', verbose_name="Владелец")
    user = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name="Студент")
    course = models.ForeignKey(Course, on_delete=models.CASCADE, verbose_name="Курс")

    flag = models.BooleanField(default=False, verbose_name="Присутствие")
    date_from = models.DateTimeField(verbose_name="Начало интервала")
    date_to = models.DateTimeField(verbose_name="Конец интервала")

    class Meta(CDEntry.Meta):
        constraints = [models.UniqueConstraint(fields=('user', 'course', 'date_from', 'date_to'),
                                               name='unique_attendance')]
        verbose_name = "Посещение"
        verbose_name_plural = "Посещаемость"

    def __str__(self):
        return f"{self.user.account} {'при' if self.flag else 'от'}сутствовал на курсе {self.course} c {self.date_from} по {self.date_to}"


"""===================================================== Filter ====================================================="""


class Filter(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name="Пользователь")
    course = models.ForeignKey(Course, on_delete=models.CASCADE, verbose_name="Курс")

    class Meta:
        verbose_name = "Фильтр"
        verbose_name_plural = "Фильтры"


"""==================================================== Contest ====================================================="""


class ContestManager(SoftDeletionManager):
    def get_new_number(self, course):
        return (self.filter(course=course).aggregate(models.Max('number')).get('number__max') or 0) + 1

    def get_queryset_for_problem(self, contest, user):
        queryset = self.get_queryset()
        return (queryset.filter(Q(id=contest.id) | Q(course__faculty=user.account.faculty, course__leaders__id=user.id))
                        .distinct())


class Contest(SoftDeletionModel, CRUDEntry):
    DEFAULT_NUMBER = 1

    course = models.ForeignKey(Course, on_delete=models.CASCADE, verbose_name="Курс")

    title = models.CharField(max_length=100, verbose_name="Заголовок")
    description = models.TextField(verbose_name="Описание", blank=True)

    number = models.PositiveSmallIntegerField(default=DEFAULT_NUMBER, verbose_name="Номер")
    hidden = models.BooleanField(default=False, verbose_name="Скрыть",
                                 help_text="Скрытый раздел отображается только тем студентам, которым назначено хотя бы"
                                           " одно задание по задаче из этого раздела")

    attachment_set = GenericRelation(Attachment, content_type_field='object_type')
    comment_set = GenericRelation(Comment, content_type_field='object_type')

    objects = ContestManager.from_queryset(SoftDeletionQuerySet)()

    class Meta(CRUDEntry.Meta):
        unique_together = ('course', 'number', 'soft_deleted')
        ordering = ('number', 'id')
        verbose_name = "Раздел"
        verbose_name_plural = "Разделы"

    @property
    def files(self):
        return [attachment.file.path for attachment in self.attachment_set.all()]

    @property
    def hidden_from_students(self):
        return self.hidden

    def visible_to(self, student):
        return not self.hidden or Assignment.objects.filter(user=student, problem__contest=self).exists()

    def get_discussion_url(self):
        return reverse('contests:contest-discussion', kwargs={'pk': self.pk})

    def __str__(self):
        return f"{self.number}. {self.title}"


"""==================================================== Problem ====================================================="""


class ProblemQuerySet(SoftDeletionQuerySet):
    def programs(self):
        return self.filter(type='Program')

    def texts(self):
        return self.filter(type='Text')

    def files(self):
        return self.filter(type='Files')

    def options(self):
        return self.filter(type='Options')

    def tests(self):
        return self.filter(type='Test')


class ProblemManager(SoftDeletionManager):
    def get_new_number(self, contest):
        return (self.filter(contest=contest).aggregate(models.Max('number')).get('number__max') or 0) + 1


class Problem(SoftDeletionModel, CRUDEntry):
    TYPE_CHOICES = (
        ('Program', "Способ ответа: программа"),
        ('Text', "Способ ответа: текст"),
        ('Verbal', "Способ ответа: устно"),
        ('Files', "Способ ответа: файлы"),
        ('Options', "Способ ответа: варианты"),
        ('Test', "Тест"),
    )

    DIFFICULTY_CHOICES = (
        (0, "Легкая"),
        (1, "Средняя"),
        (2, "Сложная"),
        (3, "Очень сложная")
    )

    LANGUAGE_CHOICES = (
        ('C++', "C++"),
        ('C', "C")
    )

    DEFAULT_DIFFICULTY = 0
    DEFAULT_LANGUAGE = 'C++'
    DEFAULT_TIME_LIMIT = 1
    DEFAULT_MEMORY_LIMIT = 64 * 1024
    DEFAULT_NUMBER = 1
    DEFAULT_SCORE_MAX = 100
    DEFAULT_SCORE_FOR_5 = 90
    DEFAULT_SCORE_FOR_4 = 75
    DEFAULT_SCORE_FOR_3 = 50

    contest = models.ForeignKey(Contest, on_delete=models.CASCADE, verbose_name="Раздел")
    sub_problems = models.ManyToManyField('self', through='SubProblem', through_fields=('problem', 'sub_problem'),
                                          symmetrical=False, verbose_name="Подзадачи")

    type = models.CharField(max_length=8, choices=TYPE_CHOICES, verbose_name="Тип")
    title = models.CharField(max_length=100, verbose_name="Заголовок")
    description = models.TextField(verbose_name="Описание", blank=True)

    number = models.PositiveSmallIntegerField(default=DEFAULT_NUMBER, verbose_name="Номер")
    score_max = models.PositiveSmallIntegerField(default=DEFAULT_SCORE_MAX, verbose_name="Максимальная оценка в баллах")
    score_for_5 = models.PositiveSmallIntegerField(default=DEFAULT_SCORE_FOR_5, verbose_name="Баллов для 5")
    score_for_4 = models.PositiveSmallIntegerField(default=DEFAULT_SCORE_FOR_4, verbose_name="Баллов для 4")
    score_for_3 = models.PositiveSmallIntegerField(default=DEFAULT_SCORE_FOR_3, verbose_name="Баллов для 3")
    difficulty = models.PositiveSmallIntegerField(choices=DIFFICULTY_CHOICES, default=DEFAULT_DIFFICULTY,
                                                  verbose_name="Сложность")
    language = models.CharField(max_length=8, choices=LANGUAGE_CHOICES, default=DEFAULT_LANGUAGE, verbose_name="Язык")
    compile_args = models.CharField(max_length=255, blank=True, verbose_name="Параметры компиляции")
    launch_args = models.CharField(max_length=255, blank=True, verbose_name="Параметры запуска")
    time_limit = models.PositiveSmallIntegerField(default=DEFAULT_TIME_LIMIT, verbose_name="Ограничение по времени")
    memory_limit = models.PositiveIntegerField(default=DEFAULT_MEMORY_LIMIT, verbose_name="Ограничение по памяти")
    is_testable = models.BooleanField(default=True, verbose_name="Проверять автоматически",
                                      help_text="Разрешить системе автоматически проверять посылки и обновлять оценку "
                                                "задания")

    attachment_set = GenericRelation(Attachment, content_type_field='object_type')
    comment_set = GenericRelation(Comment, content_type_field='object_type')

    objects = ProblemManager.from_queryset(ProblemQuerySet)()

    class Meta(CRUDEntry.Meta):
        unique_together = ('contest', 'number', 'soft_deleted')
        ordering = ('number', 'id')
        verbose_name = "Задача"
        verbose_name_plural = "Задачи"

    @property
    def course(self):
        return self.contest.course

    @property
    def files(self):
        return [attachment.file.path for attachment in self.attachment_set.all()]

    def visible_to(self, student):
        return self.contest.visible_to(student)

    def save(self, *args, **kwargs):
        if self.type not in {'Program', 'Options'}:
            self.is_testable = False
        super().save(*args, **kwargs)

    def get_score(self, submission):
        if self.type == 'Program':
            score = 2
            if submission.is_ok:
                score = 4
            elif submission.status in ['TF', 'TR', 'WA', 'NA']:
                score = 3
                if self.course.level in (6, 7):
                    passed = list(submission.execution_set.values_list('test_is_passed', flat=True).order_by())
                    if passed.count(True) not in (2, 3, 4):
                        score = 2
            return score
        if submission.score >= self.score_for_5:
            return 5
        elif submission.score >= self.score_for_4:
            return 4
        elif submission.score >= self.score_for_3:
            return 3
        return 2

    def get_latest_submission_by(self, user):
        try:
            return self.submission_set.filter(owner=user).latest('date_created')
        except ObjectDoesNotExist:
            return None

    def get_tests(self):
        return list(self.iotest_set.all()) + list(self.uttest_set.all()) + list(self.fntest_set.all())

    def run_tests(self, submission, observer, user, sandbox_type):
        state, executions = Status.UN, []
        for test in self.get_tests():
            stats = {}
            try:
                state, stats = test.run(submission, observer, self, user=user, sandbox_type=sandbox_type)
            except Exception as e:
                state, stats['exception'] = Status.EX, str(e)
            executions.append((test, stats))
            if state != Status.OK:
                break
        Execution.objects.create_set(submission, executions)
        return state

    def get_discussion_url(self):
        return reverse('contests:problem-discussion', kwargs={'pk': self.pk})

    def __str__(self):
        return f"{self.number}. {self.title}"


"""===================================================== Option ====================================================="""


class OptionQuerySet(models.QuerySet):
    def order_randomly(self):
        return self.order_by('?')


class Option(models.Model):
    problem = models.ForeignKey(Problem, on_delete=models.CASCADE, verbose_name="Задача")

    text = models.CharField(max_length=250, verbose_name="Текст")
    is_correct = models.BooleanField(default=False, verbose_name="Верный")

    objects = models.Manager.from_queryset(OptionQuerySet)()

    class Meta:
        verbose_name = "Вариант ответа"
        verbose_name_plural = "Варианты ответа"

    def __str__(self):
        return self.text


"""=================================================== SubProblem ==================================================="""


class SubProblemManager(models.Manager):
    def get_new_number(self, problem):
        return (self.filter(problem=problem).aggregate(models.Max('number')).get('number__max') or 0) + 1


class SubProblem(models.Model):
    problem = models.ForeignKey(Problem, on_delete=models.CASCADE, verbose_name="Тест")
    sub_problem = models.ForeignKey(Problem, on_delete=models.CASCADE, related_name='+', verbose_name="Задача")

    number = models.PositiveSmallIntegerField(default=Problem.DEFAULT_NUMBER, verbose_name="Номер в тесте")

    objects = SubProblemManager()

    class Meta:
        constraints = [models.UniqueConstraint(fields=('problem', 'sub_problem'), name='unique_sub_problem_in_problem'),
                       models.UniqueConstraint(fields=('problem', 'number'), name='unique_number_in_problem')]
        ordering = ('number',)
        verbose_name = "Подзадача теста"
        verbose_name_plural = "Подзадачи теста"

    def __str__(self):
        return f"{self.number}-я подзадача в тесте {self.problem.title}"


"""=============================================== SubmissionPattern ================================================"""


class SubmissionPattern(CRUDEntry):
    problems = models.ManyToManyField(Problem, related_name='submission_patterns', verbose_name="Задачи")

    title = models.CharField(max_length=100, verbose_name="Заголовок")
    description = models.TextField(verbose_name="Описание")
    pattern = models.TextField(blank=True, verbose_name="Шаблон")

    class Meta(CRUDEntry.Meta):
        verbose_name = "Шаблон посылки"
        verbose_name_plural = "Шаблоны посылок"

    def get_absolute_url(self):
        return reverse('contests:submission-pattern-detail', kwargs={'pk': self.pk})

    def __str__(self):
        return self.title


"""===================================================== Tests ======================================================"""


class BasicTest(CRUDEntry):
    problem = models.ForeignKey(Problem, on_delete=models.CASCADE, verbose_name="Задача")

    title = models.CharField(max_length=100, verbose_name="Заголовок")
    compile_args = models.CharField(max_length=255, blank=True, verbose_name="Параметры компиляции")
    compile_args_override = models.BooleanField(default=False, verbose_name="Заменить параметры компиляции из задачи?")
    launch_args = models.CharField(max_length=255, blank=True, verbose_name="Параметры запуска")
    launch_args_override = models.BooleanField(default=False, verbose_name="Заменить параметры запуска из задачи?")

    class Meta(CRUDEntry.Meta):
        abstract = True

    @property
    def course(self):
        return self.problem.contest.course

    @property
    def contest(self):
        return self.problem.contest

    def get_compile_args(self):
        compile_args = self.compile_args.split()
        if self.compile_args_override:
            return compile_args
        else:
            return self.problem.compile_args.split() + compile_args

    def get_launch_args(self):
        launch_args = self.launch_args.split()
        if self.launch_args_override:
            return launch_args
        else:
            return self.problem.launch_args.split() + launch_args

    def __str__(self):
        return self.title


class IOTest(BasicTest):
    input = models.TextField(blank=True, verbose_name="Входные данные")
    output = models.TextField(blank=True, verbose_name="Выходные данные")

    class Meta(BasicTest.Meta):
        verbose_name = "IO-тест"
        verbose_name_plural = "IO-тесты"

    def run(self, submission, observer, _, user=None, sandbox_type='subprocess'):
        sources = submission.files
        workdir = os.path.dirname(sources[0])
        state, stats = Status.UN, {}
        Sandbox = get_sandbox_class(sandbox_type)
        with Sandbox(workdir) as sandbox:
            executable = sandbox.path('exe')
            observer.set_progress('Компилируем', 5, 100)
            state = sandbox.compile(sources, executable, language=self.problem.language, args=self.get_compile_args(),
                                    timeout=self.problem.time_limit, stats=stats)
            if state == Status.OK:
                input_file = sandbox.create('input.txt', self.input)
                output_file = sandbox.path('output.txt')
                observer.set_progress('Проверяем', 60, 100)
                state = sandbox.execute(executable, args=[input_file, output_file], timeout=self.problem.time_limit,
                                        stats=stats)
                if state == Status.OK:
                    output = sandbox.read(output_file)
                    if output is None:
                        state = Status.NA
                    elif diff(output, self.output):
                        state = Status.WA
                        stats['test_output'] = output
                    else:
                        stats['test_is_passed'] = True
        return state, stats


class UTTest(BasicTest):
    attachment_set = GenericRelation(Attachment, content_type_field='object_type')

    class Meta(BasicTest.Meta):
        verbose_name = "UT-тест"
        verbose_name_plural = "UT-тесты"

    @property
    def files(self):
        return [attachment.file.path for attachment in self.attachment_set.all()]

    def run(self, submission, observer, _, user=None, sandbox_type='subprocess'):
        sources = submission.files
        workdir = os.path.dirname(sources[0])
        state, stats = Status.UN, {}
        Sandbox = get_sandbox_class(sandbox_type)
        with Sandbox(workdir) as sandbox:
            test_sources = sandbox.fetch(self.files)
            executable = sandbox.path('exe')
            observer.set_progress('Компилируем', 5, 100)
            state = sandbox.compile(sources + test_sources, executable, language=self.problem.language,
                                    args=self.get_compile_args(), timeout=self.problem.time_limit, stats=stats)
            if state == Status.OK:
                observer.set_progress('Проверяем', 60, 100)
                state = sandbox.execute(executable, args=self.get_launch_args(), timeout=self.problem.time_limit,
                                        stats=stats)
                if state == Status.OK:
                    returncode = stats.get('execution_returncode')
                    if returncode != 0:
                        state = Status.TF
                    else:
                        stats['test_is_passed'] = True
        return state, stats


class FNTest(CRUDEntry):
    HANDLER_CHOICES = (
        ('fsm.init', "файловый менеджер: тривиальный"),
        ('fsm.comm', "файловый менеджер: общий"),
        ('fsm.spec', "файловый менеджер: специальный"),
        ('mem.init', "менеджер памяти: тривиальный"),
        ('mem.comm', "менеджер памяти: общий"),
        ('mem.spec', "менеджер памяти: специальный"),
        ('lss.initial', "линейные системы: тривиальный"),
        ('lss.main', "линейные системы: основной"),
        ('lss.worst_lu', "линейные системы: worst lu"),
        ('lss.worst_lu_t', "линейные системы: worst lu t"),
        ('lss.large', "линейные системы: большие размерности"),
        ('lss.degenerate', "линейные системы: вырожденная матрица"),
        ('lss.insoluble', "линейные системы: несовместная матрица"),
        ('lss.degenerate_major_minor', "линейные системы: вырожденный минор"),
        ('evc.init', "собственные значения: тривиальный"),
        ('evc.main', "собственные значения: основной"),
        ('evc.extd', "собственные значения: расширенный"),
    )  # catch choices from problems.modules

    problems = models.ManyToManyField(Problem, verbose_name="Задачи")

    title = models.CharField(max_length=100, verbose_name="Заголовок")
    handler = models.CharField(max_length=100, choices=HANDLER_CHOICES, verbose_name="Обработчик")

    class Meta(CRUDEntry.Meta):
        verbose_name = "FN-тест"
        verbose_name_plural = "FN-тесты"

    def run(self, submission, observer, problem, user=None, sandbox_type='subprocess'):
        module_name, function_name = self.handler.split('.')
        module = importlib.import_module('tools.problems.' + module_name)
        return getattr(module, function_name)(submission, observer, problem, self, sandbox_type=sandbox_type)

    def __str__(self):
        return self.title


"""=================================================== Assignment ==================================================="""


class AssignmentQuerySet(models.QuerySet):
    def for_course_table(self, course, students):
        latest_submission = (Submission.objects.filter(owner_id=models.OuterRef('user_id'),
                                                       problem_id=models.OuterRef('problem_id'))
                             .order_by('-date_created')[:1])
        return (self.filter(user__in=students.values_list('user'), problem__contest__course=course)
                .select_related('user', 'problem')
                .annotate(latest_submission_status=models.Subquery(latest_submission.values_list('status')))
                .annotate(latest_submission_date_created=models.Subquery(latest_submission.values_list('date_created')))
                .order_by('user__account', 'problem__contest', 'date_created', 'problem__number'))

    def rollback_score(self):
        return self.update(score=models.F('score') - 1)

    def to_rollback(self, submissions):
        return self.filter(id__in=submissions.values_list('assignment_id', flat=True), score__gt=3)

    def progress(self):
        naccomplished = self.filter(score__in=[5, 4]).count()
        ntotal = self.count() or 1
        return naccomplished * 100 // ntotal


class AssignmentManager(models.Manager):
    def create_random_set(self, owner, contest, type, limit_per_user, submission_limit, deadline, params):
        """ create random set of assignments with problems of given contest
            aligning their number to limit_per_user for contest.course.level students """
        problem_ids = contest.problem_set.filter(type=type).order_by('number').values_list('id', flat=True)
        student_ids = Account.students.apply_common_filters(params).values_list('user_id', flat=True)
        new_assignments = []
        for student_id in student_ids:
            assigned_problem_ids = (self.filter(user_id=student_id, problem_id__in=problem_ids)
                                    .select_related('problem')
                                    .values_list('problem_id', flat=True))
            problem_id_set = set(problem_ids)
            assigned_problem_id_set = set(assigned_problem_ids)
            if len(assigned_problem_id_set) < limit_per_user:
                to_assign_problem_id_list = []
                unassigned_problem_id_list = list(problem_id_set - assigned_problem_id_set)
                while unassigned_problem_id_list and len(assigned_problem_id_set) < limit_per_user:
                    i = random.randint(0, len(unassigned_problem_id_list) - 1)
                    problem_id = unassigned_problem_id_list.pop(i)
                    to_assign_problem_id_list.append(problem_id)
                    assigned_problem_id_set.add(problem_id)
                to_assign_problem_id_list = list(filter(lambda x: x in to_assign_problem_id_list, problem_ids))
                for to_assign_problem_id in to_assign_problem_id_list:
                    new_assignment = Assignment(owner_id=owner.id, user_id=student_id, problem_id=to_assign_problem_id,
                                                submission_limit=submission_limit)
                    if deadline is not None:
                        new_assignment.deadline = deadline
                    new_assignments.append(new_assignment)
        Assignment.objects.bulk_create(new_assignments)


class Assignment(CRUDEntry):
    DEFAULT_SCORE = 0
    DEFAULT_SCORE_MAX = 5
    DEFAULT_SUBMISSION_LIMIT = 10

    owner = models.ForeignKey(User, on_delete=models.CASCADE, related_name='+', verbose_name="Владелец")
    user = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name="Студент")
    problem = models.ForeignKey(Problem, on_delete=models.CASCADE, verbose_name="Задача")

    score = models.PositiveSmallIntegerField(default=DEFAULT_SCORE,
                                             validators=[MinValueValidator(0), MaxValueValidator(5)],
                                             verbose_name="Оценка")
    score_max = models.PositiveSmallIntegerField(default=DEFAULT_SCORE_MAX,
                                                 validators=[MinValueValidator(3), MaxValueValidator(5)],
                                                 verbose_name="Максимальная оценка")
    score_is_locked = models.BooleanField(default=False, verbose_name="Заблокировать оценку",
                                          help_text="Заблокированная оценка не может быть изменена системой "
                                                    "автоматической проверки посылок")
    submission_limit = models.PositiveSmallIntegerField(default=DEFAULT_SUBMISSION_LIMIT,
                                                        verbose_name="Ограничение количества посылок")
    remark = models.CharField(max_length=255, blank=True, verbose_name="Пометка", help_text="Для преподавателей")
    deadline = models.DateTimeField(null=True, blank=True, verbose_name="Принимать посылки до")

    attachment_set = GenericRelation(Attachment, content_type_field='object_type')
    comment_set = GenericRelation(Comment, content_type_field='object_type')

    objects = AssignmentManager.from_queryset(AssignmentQuerySet)()

    class Meta(CRUDEntry.Meta):
        unique_together = ('user', 'problem')
        verbose_name = "Задание"
        verbose_name_plural = "Задания"

    @property
    def credit_incomplete(self):
        return self.credit is not None and self.credit.score < 3

    @property
    def credit(self):
        try:
            return Credit.objects.get(user=self.user, course=self.course)
        except Credit.DoesNotExist:
            return None

    @property
    def course(self):
        return self.problem.contest.course

    @property
    def contest(self):
        return self.problem.contest

    @property
    def hidden_from_students(self):
        return self.deadline is not None and timezone.now() < self.deadline

    def get_latest_submission(self):
        return self.problem.get_latest_submission_by(self.user)

    def update_score(self, submission):
        if not self.score_is_locked:
            score = self.problem.get_score(submission)
            if self.score < score:
                self.score = min(score, self.score_max)
                self.save(update_fields=['score'])

    def get_discussion_url(self):
        return reverse('contests:assignment-discussion', kwargs={'pk': self.pk})

    def __str__(self):
        return f"Задание для {self.user.account.get_short_name()}: {self.problem}"


@receiver(models.signals.post_save, sender=Assignment)
def link_existing_submissions(sender, instance, created, **kwargs):
    """ links existing submissions to newly created `Assignment` object. """
    if created:
        submissions = Submission.objects.filter(assignment__isnull=True, owner=instance.user, problem=instance.problem)
        submissions.update(assignment=instance)


"""=================================================== Submission ==================================================="""


class SubmissionQuerySet(models.QuerySet):
    def rollback_status(self):
        return self.update(status='TR')

    def to_rollback(self, problem_id):
        return self.filter(problem_id=problem_id, status__in=['OK', 'TR']).filter(models.Exists(
            Credit.objects.filter(user_id=models.OuterRef('owner_id'),
                                  course_id=models.OuterRef('problem__contest__course_id'),
                                  score=Credit.DEFAULT_SCORE)
        ))


class SubmissionManager(models.Manager):
    def backup(self, submission_list):
        stream = io.BytesIO()
        with zipfile.ZipFile(stream, 'w') as zip_file:
            for submission in submission_list:
                for f in submission.files:
                    try:
                        _, filename = os.path.split(f)
                        zip_path = "{contest_id}/{problem_num}/{user_id}/{submission_id}/{filename}".format(
                            contest_id=submission.problem.contest_id,
                            problem_num=submission.problem.number,
                            user_id=submission.owner_id,
                            submission_id=submission.id,
                            filename=filename
                        )
                        zip_file.write(f, zip_path)
                    except FileNotFoundError:
                        pass
        return stream.getvalue()


class Submission(CRDEntry):
    STATUS_CHOICES = (
        ('OK', "Задача решена"),
        ('PS', "Задача решена частично"),
        ('TF', "Тест провален"),
        ('TR', "Требуется проверка"),
        ('WA', "Неверный ответ"),
        ('NA', "Ответ отсутствует"),
        ('TL', "Превышено ограничение по времени"),
        ('ML', "Превышено ограничение по памяти"),
        ('CL', "Превышено ограничение по времени компиляции"),
        ('FE', "Ошибка операции с плавающей точкой"),
        ('SF', "Ошибка при работе с памятью"),
        ('RE', "Ошибка выполнения"),
        ('CE', "Ошибка компиляции"),
        ('UE', "Ошибка кодировки"),
        ('PE', "Ошибка комплектации"),
        ('EX', "Неизвестная ошибка"),
        ('EV', "Посылка проверяется"),
        ('UN', "Посылка не проверена")
    )
    DEFAULT_STATUS = 'UN'
    DEFAULT_SCORE = 0

    problem = models.ForeignKey(Problem, on_delete=models.CASCADE, related_query_name="submission",
                                verbose_name="Задача")
    assignment = models.ForeignKey(Assignment, on_delete=models.SET_NULL, null=True, verbose_name="Задание")
    main_submission = models.ForeignKey('self', on_delete=models.CASCADE, null=True, related_name="sub_submissions",
                                        verbose_name="Подпосылки")
    options = models.ManyToManyField(Option, verbose_name="Варианты ответа")

    footprint = models.TextField()
    status = models.CharField(max_length=2, choices=STATUS_CHOICES, default=DEFAULT_STATUS, verbose_name="Статус")
    score = models.PositiveSmallIntegerField(default=DEFAULT_SCORE, verbose_name="Оценка в баллах")
    text = models.TextField(null=True, blank=True, verbose_name="Ответ")
    task_id = models.UUIDField(null=True, blank=True, verbose_name="Идентификатор асинхронной задачи")
    moss_to_submissions = models.CharField(max_length=200, null=True,
                                           validators=[validate_comma_separated_integer_list],
                                           verbose_name="С посылками MOSS")
    moss_report_url = models.URLField(null=True, verbose_name="Ссылка на отчет MOSS")

    attachment_set = GenericRelation(Attachment, content_type_field='object_type')
    comment_set = GenericRelation(Comment, content_type_field='object_type')

    objects = SubmissionManager.from_queryset(SubmissionQuerySet)()

    class Meta(CRDEntry.Meta):
        permissions = [
            ("evaluate_submission", "Проверять Посылку"),
            ("download_submission", "Скачивать Посылку"),
            ("moss_submission", "Отправлять на проверку в MOSS"),
        ]
        ordering = ('-date_created',)
        verbose_name = "Посылка"
        verbose_name_plural = "Посылки"

    @property
    def course(self):
        return self.problem.contest.course

    @property
    def contest(self):
        return self.problem.contest

    @property
    def short_title(self):
        return "П{}осылка {}".format("" if self.main_submission is None else "одп", self.id)

    @property
    def hidden_from_students(self):
        assignment = self.get_assignment()
        return assignment is not None and assignment.hidden_from_students

    def get_assignment(self):
        return self.assignment if self.main_submission is None else self.main_submission.assignment

    @property
    def is_ok(self):
        return self.status == 'OK'

    @property
    def is_un(self):
        return self.status == 'UN'

    def get_score_percentage(self):
        return self.score * 100 // self.problem.score_max

    def moss_to_submissions_list(self):
        return self.moss_to_submissions.split(',')

    @property
    def files(self):
        return [attachment.file.path for attachment in self.attachment_set.all()]

    @property
    def has_footprint_increments(self):
        if self.problem.type != 'Text' or not self.footprint:
            return False
        footprint = json.loads(self.footprint)
        if isinstance(footprint, list):
            prev = 0
            for cur in footprint:
                if cur - prev > 50:
                    return True
                prev = cur
        return False

    def get_files_as_zip(self):
        stream = io.BytesIO()
        zip_dirname = str(self.pk)
        with zipfile.ZipFile(stream, 'w') as zip_file:
            for file in self.files:
                _, filename = os.path.split(file)
                zip_path = os.path.join(zip_dirname, filename)
                zip_file.write(file, zip_path)
        return stream.getvalue()

    def inspect(self, observer):
        return Status.OK

    def test(self, observer, user, sandbox_type):
        return self.problem.run_tests(self, observer, user, sandbox_type)

    def update(self, state):
        self.task_id = None
        self.status = str(state)
        self.save()

    def evaluate(self, observer, user, sandbox_type):
        state = self.inspect(observer)
        if state == Status.OK:
            state = self.test(observer, user, sandbox_type)
        self.update(state)

    def update_assignment(self):
        if self.assignment is not None:
            self.assignment.update_score(self)

    def update_main_score(self):
        max_score = sum(self.problem.sub_problems.values_list('score_max', flat=True))
        score_sum = sum(self.sub_submissions.values_list('score', flat=True))
        self.score = score_sum * self.problem.score_max // max_score

    def update_main_status(self):
        if self.is_un:
            statuses = list(choice for choice, _ in self.STATUS_CHOICES)
            acquired_statuses = set(self.sub_submissions.values_list('status', flat=True))
            for status in reversed(statuses):
                if status in acquired_statuses:
                    self.status = status
                    break
        if self.is_un:
            score = self.problem.get_score(self)
            if score == 5:
                self.status = 'OK'
            elif score == 4:
                self.status = 'PS'
            elif score <= 3:
                self.status = 'WA'

    def evaluate_options(self):
        correct_option_ids = set(self.problem.option_set.filter(is_correct=True).values_list('id', flat=True))
        chosen_option_ids = set(self.options.values_list('id', flat=True))
        if chosen_option_ids == correct_option_ids:
            self.score = self.problem.score_max
            self.status = 'OK'
        else:
            self.score = 0
            self.status = 'WA'

    def get_discussion_url(self):
        return self.get_absolute_url()

    def save(self, *args, **kwargs):
        if self.footprint is None:  # TODO: remove if not needed in frontend
            self.footprint = "[]"
        created = self._state.adding
        super().save(*args, **kwargs)
        if created and self.problem.type == 'Options':
            self.evaluate_options()
            super().save(update_fields=['status', 'score'])
        if self.main_submission is not None:
            self.main_submission.update_main_score()
            self.main_submission.update_main_status()
            self.main_submission.save(update_fields=['status', 'score'])
        elif created:
            course_leaders = self.course.leaders.filter(Q(account__faculty=self.owner.account.faculty) |
                                                        Q(account__faculty__short_name="МФК"))
            course_leaders = course_leaders.values_list('id', flat=True)
            Notification.objects.notify(course_leaders, subject=self.owner, action="отправил посылку", object=self,
                                        relation="к задаче", reference=self.problem)
        self.update_assignment()

    def __str__(self):
        return f"Посылка от {self.owner.account.get_short_name()} к задаче {self.problem}"


"""=================================================== Execution ===================================================="""


class ExecutionManager(models.Manager):
    def create_set(self, submission, executions):
        self.filter(submission=submission).delete()
        new_executions = []
        for test, stats in executions:
            if 'date_created' not in stats:
                stats['date_created'] = timezone.now()
            elif timezone.is_naive(stats['date_created']):
                stats['date_created'] = timezone.make_aware(stats['date_created'])
            new_execution = Execution(submission=submission, test_type=ContentType.objects.get_for_model(type(test)),
                                      test_id=test.id, **stats)
            new_executions.append(new_execution)
        return self.bulk_create(new_executions)


class Execution(models.Model):
    submission = models.ForeignKey(Submission, on_delete=models.CASCADE, verbose_name="Посылка")
    test_type = models.ForeignKey(ContentType, on_delete=models.CASCADE)
    test_id = models.PositiveIntegerField()
    test = GenericForeignKey(ct_field='test_type', fk_field='test_id')

    compilation_time = models.FloatField(default=0.0, verbose_name="Время компиляции")
    compilation_command = models.TextField(default="", verbose_name="Команда компиляции")
    compilation_stdout = models.TextField(default="", verbose_name="Вывод компилятора в stdout")
    compilation_stderr = models.TextField(default="", verbose_name="Вывод компилятора в stderr")

    execution_memory = models.PositiveIntegerField(default=0, verbose_name="Использовано памяти")
    execution_time = models.FloatField(default=0.0, verbose_name="Время выполнения")
    execution_command = models.TextField(default="", verbose_name="Команда выполнения")
    execution_returncode = models.IntegerField(default=0, verbose_name="Код возврата")
    execution_stdout = models.TextField(default="", verbose_name="Вывод программы в stdout")
    execution_stderr = models.TextField(default="", verbose_name="Вывод программы в stderr")

    test_is_passed = models.BooleanField(default=False, verbose_name="Тест пройден?")
    test_input = models.TextField(default="", verbose_name="Входные данные")
    test_output = models.TextField(default="", verbose_name="Выходные данные")
    test_output_correct = models.TextField(default="", verbose_name="Верные выходные данные")
    test_summary = models.TextField(default="", verbose_name="Сводка")

    exception = models.TextField(default="", verbose_name="Исключение")

    date_created = models.DateTimeField(verbose_name="Дата создания")

    objects = ExecutionManager()

    class Meta:
        ordering = ('date_created',)
        verbose_name = "Запуск"
        verbose_name_plural = "Запуски"

    def __str__(self):
        return f"Запуск: {self.submission} на тесте {self.test} от {self.date_created}"
